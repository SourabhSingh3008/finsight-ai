from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_business_doc():
    doc = Document()
    
    # Custom styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    title = doc.add_heading('FinSight AI', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph('Product Requirements & Business Description Document')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # 1. Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        "FinSight AI is an intelligent, client-side personal finance web application designed to empower individuals "
        "by instantly transforming unstructured banking data into actionable financial insights. Leveraging the advanced "
        "capabilities of generative artificial intelligence (Google Gemini 1.5 Flash), FinSight AI automates expense "
        "categorization, detects hidden recurring subscriptions, and acts as a localized, secure, and highly personalized "
        "financial advisor."
    )

    # 2. Problem Statement
    doc.add_heading('2. Problem Statement & Market Opportunity', level=1)
    doc.add_paragraph(
        "The modern consumer's financial landscape is fragmented. Bank statements are notoriously difficult to read, "
        "often filled with cryptic merchant codes and unstructured text. Consequently, consumers struggle to:"
    )
    doc.add_paragraph("• Accurately track categorized expenditures month-over-month.", style='List Bullet')
    doc.add_paragraph("• Identify and cancel unused or hidden recurring subscriptions (the 'subscription trap').", style='List Bullet')
    doc.add_paragraph("• Obtain personalized, actionable financial advice without paying high fees for financial advisors.", style='List Bullet')
    doc.add_paragraph("• Trust third-party budgeting apps with sensitive banking credentials (via Plaid or similar services), raising significant data privacy concerns.", style='List Bullet')

    # 3. Value Proposition
    doc.add_heading('3. The FinSight AI Value Proposition', level=1)
    doc.add_paragraph(
        "FinSight AI provides a frictionless and secure alternative to traditional budgeting applications. By dropping "
        "the requirement for direct bank integration, users can simply paste their raw transaction history (from a PDF "
        "or web portal) into the application. The on-device, client-side architecture immediately processes this data "
        "using a user-provided LLM API key, guaranteeing that personal financial data is never stored on a centralized backend."
    )

    # 4. Target Audience
    doc.add_heading('4. Target Audience', level=1)
    doc.add_paragraph("• Young Professionals & Millennials: Seeking modern, aesthetic tools to manage multiple subscriptions and discretionary spending.", style='List Bullet')
    doc.add_paragraph("• Privacy-Conscious Users: Individuals hesitant to link their actual bank accounts to third-party aggregation apps.", style='List Bullet')
    doc.add_paragraph("• Gig Workers & Freelancers: Needing quick, ad-hoc categorization of messy income and expense streams without complex accounting software.", style='List Bullet')

    # 5. Core Features & Product Capabilities
    doc.add_heading('5. Core Features', level=1)
    p = doc.add_paragraph()
    p.add_run("5.1 Seamless Unstructured Data Ingestion\n").bold = True
    p.add_run("Users can copy and paste raw text arrays from any bank statement. The built-in AI natively understands dates, amounts, and descriptions without requiring a standardized CSV format.")
    
    p = doc.add_paragraph()
    p.add_run("5.2 Intelligent Categorization Engine\n").bold = True
    p.add_run("Instead of relying on rigid, hardcoded keyword matching, the AI contextually understands merchants (e.g., categorizing 'Uber Eats' as Dining and 'Uber' as Transportation).")
    
    p = doc.add_paragraph()
    p.add_run("5.3 Subscription Detection & Alerting\n").bold = True
    p.add_run("The system isolates recurring charges, aggregates the total monthly subscription burden, and flags them for the user, helping to mitigate subscription fatigue.")
    
    p = doc.add_paragraph()
    p.add_run("5.4 The AI Financial Advisor\n").bold = True
    p.add_run("Beyond just dashboards, the app generates three immediate, contextual insights. For example, warning a user about high utility bills or suggesting a cut in entertainment spending based on the specific dataset provided.")

    # 6. Technical Architecture & AI Integration
    doc.add_heading('6. Technical Architecture & AI Integration', level=1)
    doc.add_paragraph(
        "FinSight AI is built on a lightweight, highly responsive pure frontend stack (HTML5, CSS3, Vanilla JS). "
        "This completely eliminates server costs and backend maintenance. "
    )
    doc.add_paragraph(
        "AI Integration Strategy:\n"
        "• Model: Google Gemini 1.5 Flash (chosen for its speed and high accuracy with unstructured text). \n"
        "• Prompt Engineering: The system uses deterministic prompt framing, forcing the AI to return a strict, pre-defined JSON schema containing 'total_spend', 'categories', 'subscriptions', and 'insights'.\n"
        "• Security: The user inputs their own API key, which is saved locally in the browser's LocalStorage. No data touches a FinSight-owned database."
    )

    # 7. User Journey
    doc.add_heading('7. User Journey', level=1)
    doc.add_paragraph("1. Onboarding: User lands on a premium, glassmorphic UI and inputs their Gemini API key.", style='List Number')
    doc.add_paragraph("2. Data Input: User pastes raw transaction strings into a frictionless text area.", style='List Number')
    doc.add_paragraph("3. Analysis: The app displays a micro-animation loading state while the LLM parses the context.", style='List Number')
    doc.add_paragraph("4. Dashboard: A dynamic financial health dashboard is rendered, summarizing total spend, charting category data, listing subscriptions, and providing narrative AI advice.", style='List Number')

    # 8. Future Roadmap & Monetization
    doc.add_heading('8. Future Roadmap & Monetization', level=1)
    doc.add_paragraph(
        "While the current MVP is a client-side utility, Phase 2 will introduce:"
    )
    doc.add_paragraph("• Direct PDF/Image upload leveraging Gemini Vision for OCR.", style='List Bullet')
    doc.add_paragraph("• Plaid Integration for users who opt-in for automated syncing.", style='List Bullet')
    doc.add_paragraph("• Export to Excel/CSV for accounting purposes.", style='List Bullet')
    doc.add_paragraph(
        "Monetization Strategy: Transitioning to a Freemium model where basic text pasting is free (using a heavily rate-limited native key), "
        "but pro features like OCR, automated bank syncing, and multi-month historical trend analysis require a $4.99/mo subscription."
    )

    doc.save('FinSight_AI_Business_Document.docx')
    print("Business DOCX created successfully.")

if __name__ == '__main__':
    create_business_doc()
